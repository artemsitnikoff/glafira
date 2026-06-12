"""Парсинг резюме и автозаполнение полей кандидата"""

import asyncio
import logging
import re
from io import BytesIO
from pathlib import Path
from uuid import UUID

from docx import Document
from sqlalchemy import select
from sqlalchemy.orm import selectinload
from sqlalchemy.ext.asyncio import AsyncSession

from .client import call_json
from .prompts import RESUME_PARSE_PROMPT
from ...models import Candidate, CandidateExperience, CandidateSkill, CandidateEducation

logger = logging.getLogger(__name__)


def _to_int(value) -> int | None:
    """Безопасно привести ответ LLM к int. LLM по промпту возвращает зарплату строкой
    ('180 000 ₽ на руки'), а колонка — INTEGER. Берём первое число, чистим разделители."""
    if value is None or isinstance(value, bool):
        return None
    if isinstance(value, int):
        n = value
    elif isinstance(value, float):
        n = int(value)
    elif isinstance(value, str):
        m = re.search(r'\d[\d\s]*\d|\d', value)
        if not m:
            return None
        try:
            n = int(re.sub(r'\s', '', m.group()))
        except ValueError:
            return None
    else:
        return None
    # Санити: отрицательные/абсурдные отбросить (PG INTEGER до ~2.1 млрд)
    if n < 0 or n > 1_000_000_000:
        return None
    return n


def _to_str(value, maxlen: int) -> str | None:
    """Привести к строке нужной длины (обрезать под лимит колонки), пустое/None — в None."""
    if value is None:
        return None
    s = str(value).strip()
    return s[:maxlen] if s else None


def _extract_pdf_text(content: bytes) -> str | None:
    """Синхронный парсинг PDF (выносится в asyncio.to_thread, чтобы не блокировать event loop
    на время извлечения текста — на больших/многостраничных резюме это секунды)."""
    from pypdf import PdfReader
    reader = PdfReader(BytesIO(content))
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    return "\n".join(text_parts) if text_parts else None


def _extract_docx_text(content: bytes) -> str | None:
    """Синхронный парсинг .docx через python-docx (также выносится в to_thread)."""
    try:
        doc = Document(BytesIO(content))
        text_parts = []

        # Извлекаем текст из параграфов
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text_parts.append(cell_text)

        return "\n".join(text_parts) if text_parts else None
    except Exception:
        return None


async def extract_resume_text(content: bytes, filename: str) -> str | None:
    """Extract text content from resume file"""
    file_ext = Path(filename).suffix.lower()

    if file_ext in {'.txt', '.md'}:
        try:
            text = content.decode('utf-8', errors='ignore')
            logger.info(f"Extracted {len(text)} characters from text file {filename}")
            return text if text.strip() else None
        except Exception:
            return None

    elif file_ext == '.pdf':
        try:
            text = await asyncio.to_thread(_extract_pdf_text, content)
            if text:
                logger.info(f"Extracted {len(text)} characters from PDF {filename}")
            else:
                logger.warning(f"No text extracted from PDF {filename} (возможно скан/картинка)")
            return text
        except Exception as e:
            logger.warning(f"Failed to extract PDF text: {e}")
            return None

    elif file_ext == '.docx':
        try:
            text = await asyncio.to_thread(_extract_docx_text, content)
            if text:
                logger.info(f"Extracted {len(text)} characters from DOCX {filename}")
            else:
                logger.warning(f"No text extracted from DOCX {filename}")
            return text
        except Exception as e:
            logger.warning(f"Failed to extract DOCX text: {e}")
            return None

    elif file_ext == '.doc':
        # .doc (старый бинарный формат) не поддерживается честно
        logger.info(f"DOC format not supported: {filename}")
        return None

    return None


async def parse_resume_to_dict(content: bytes, filename: str) -> dict | None:
    """Parse resume to structured dict without saving to DB.

    Returns:
        dict with resume fields or None if format not supported / no text extracted
    """
    # Extract text from file
    text = await extract_resume_text(content, filename)
    if not text:
        logger.info(f"No text extracted from {filename} for parse_resume_to_dict")
        return None

    try:
        # Call Claude to parse structured data
        parsed_data = await call_json(
            system=RESUME_PARSE_PROMPT,
            user=text,
            max_tokens=16000
        )

        # Коэрсинг полей к нужным типам и длинам
        result = {
            # ФИО (новые поля)
            "first_name": _to_str(parsed_data.get("first_name"), 120),
            "last_name": _to_str(parsed_data.get("last_name"), 120),
            "middle_name": _to_str(parsed_data.get("middle_name"), 120),

            # Базовые скаляры
            "phone": _to_str(parsed_data.get("phone"), 20),
            "email": _to_str(parsed_data.get("email"), 255),
            "city": _to_str(parsed_data.get("city"), 120),
            "salary_expectation": _to_int(parsed_data.get("salary_expectation")),
            "last_position": _to_str(parsed_data.get("last_position"), 255),
            "last_company": _to_str(parsed_data.get("last_company"), 255),
            "last_period": _to_str(parsed_data.get("last_period"), 120),
            "about": _to_str(parsed_data.get("about"), 5000),

            # Структурированные секции
            "experience": [],
            "skills": [],
            "education": [],
            "languages": []
        }

        # Опыт работы - только записи с непустым position
        if parsed_data.get("experience"):
            for exp_data in parsed_data["experience"]:
                if position := _to_str(exp_data.get("position"), 255):
                    result["experience"].append({
                        "position": position,
                        "company": _to_str(exp_data.get("company"), 255),
                        "period": _to_str(exp_data.get("period"), 120),
                        "description": _to_str(exp_data.get("description"), 10000)
                    })

        # Навыки - строки
        if parsed_data.get("skills"):
            for skill_data in parsed_data["skills"]:
                if skill := _to_str(skill_data, 120):
                    result["skills"].append(skill)

        # Образование
        if parsed_data.get("education"):
            for edu_data in parsed_data["education"]:
                result["education"].append({
                    "institution": _to_str(edu_data.get("institution"), 255),
                    "specialty": _to_str(edu_data.get("specialty"), 255),
                    "years": _to_str(edu_data.get("years"), 40)
                })

        # Языки
        if parsed_data.get("languages"):
            for lang_data in parsed_data["languages"]:
                if lang := _to_str(lang_data, 120):
                    result["languages"].append(lang)

        return result

    except Exception as e:
        logger.warning(f"Resume parsing failed for {filename}: {e}")
        return None


async def parse_and_apply_resume(
    session: AsyncSession,
    *,
    candidate_id: UUID,
    content: bytes,
    filename: str,
    company_id: UUID
) -> None:
    """Parse resume and update candidate fields"""
    # Parse resume to structured dict
    parsed_data = await parse_resume_to_dict(content, filename)
    if not parsed_data:
        logger.info(f"No data extracted from {filename}")
        return

    # Get candidate с eager-загрузкой связей — иначе доступ к candidate.experience/skills/
    # education в async-режиме триггерит ленивый SELECT вне greenlet → "greenlet_spawn has
    # not been called" и парс падает.
    result = await session.execute(
        select(Candidate)
        .options(
            selectinload(Candidate.experience),
            selectinload(Candidate.skills),
            selectinload(Candidate.education),
        )
        .where(
            Candidate.id == candidate_id,
            Candidate.company_id == company_id,
            Candidate.deleted_at.is_(None)
        )
    )
    candidate = result.scalar_one_or_none()
    if not candidate:
        logger.warning(f"Candidate {candidate_id} not found")
        return

    # Update resume_text if it was None (извлекаем текст заново для сохранения)
    if candidate.resume_text is None:
        text = await extract_resume_text(content, filename)
        if text:
            candidate.resume_text = text

    try:
        # Заполняем только пустые поля кандидата скалярными значениями
        # (данные уже обработаны в parse_resume_to_dict)

        # ФИО - заполняем только пустые поля
        if candidate.first_name is None and parsed_data.get("first_name"):
            candidate.first_name = parsed_data["first_name"]

        if candidate.last_name is None and parsed_data.get("last_name"):
            candidate.last_name = parsed_data["last_name"]

        if candidate.middle_name is None and parsed_data.get("middle_name"):
            candidate.middle_name = parsed_data["middle_name"]

        # Базовые скаляры — только если поле пустое
        if candidate.last_position is None and parsed_data.get("last_position"):
            candidate.last_position = parsed_data["last_position"]

        if candidate.last_company is None and parsed_data.get("last_company"):
            candidate.last_company = parsed_data["last_company"]

        if candidate.last_period is None and parsed_data.get("last_period"):
            candidate.last_period = parsed_data["last_period"]

        if candidate.salary_expectation is None and parsed_data.get("salary_expectation") is not None:
            candidate.salary_expectation = parsed_data["salary_expectation"]

        if candidate.city is None and parsed_data.get("city"):
            candidate.city = parsed_data["city"]

        if candidate.phone is None and parsed_data.get("phone"):
            candidate.phone = parsed_data["phone"]

        if candidate.email is None and parsed_data.get("email"):
            candidate.email = parsed_data["email"]

        # «Обо мне» — самоописание кандидата (раздел резюме) → resume_summary
        if candidate.resume_summary is None and parsed_data.get("about"):
            candidate.resume_summary = parsed_data["about"]

        # Структурные записи — создаём только если у кандидата их ещё НЕТ (не затираем ручные правки)

        # Опыт работы
        if not candidate.experience and parsed_data.get("experience"):
            exp_count = 0
            created_exp = []
            for idx, exp_data in enumerate(parsed_data["experience"]):
                # position уже проверен в parse_resume_to_dict
                experience = CandidateExperience(
                    candidate_id=candidate.id,
                    company_id=company_id,
                    position=exp_data["position"],
                    company=exp_data["company"],
                    period=exp_data["period"],
                    description=exp_data["description"],
                    order_index=idx
                )
                session.add(experience)
                created_exp.append(experience)
                exp_count += 1

            # Синхронизируем денормализованные «последнее место» со самой свежей записью опыта,
            # чтобы мета карточки совпадала с опытом (а не с устаревшими/ручными last_*).
            from ...services.candidate import pick_latest_experience
            latest = pick_latest_experience(created_exp)
            if latest:
                candidate.last_position = latest.position
                candidate.last_company = latest.company
                candidate.last_period = latest.period

            logger.info(f"Created {exp_count} experience records for candidate {candidate_id}")

        # Навыки
        if not candidate.skills and parsed_data.get("skills"):
            skill_count = 0
            for idx, skill in enumerate(parsed_data["skills"]):
                candidate_skill = CandidateSkill(
                    candidate_id=candidate.id,
                    company_id=company_id,
                    skill=skill,
                    order_index=idx
                )
                session.add(candidate_skill)
                skill_count += 1

            logger.info(f"Created {skill_count} skill records for candidate {candidate_id}")

        # Образование
        if not candidate.education and parsed_data.get("education"):
            edu_count = 0
            for idx, edu_data in enumerate(parsed_data["education"]):
                education = CandidateEducation(
                    candidate_id=candidate.id,
                    company_id=company_id,
                    institution=edu_data["institution"],
                    specialty=edu_data["specialty"],
                    years=edu_data["years"],
                    order_index=idx
                )
                session.add(education)
                edu_count += 1

            logger.info(f"Created {edu_count} education records for candidate {candidate_id}")

        # Доп. поля для блока «Дополнительно» карточки (languages/relocation/business_trips/remote)
        # → пишем в candidate.extra (JSONB), как resume_gen. Реассайн обязателен (иначе SQLAlchemy
        # не заметит мутацию dict). Не затираем уже заполненные ключи.
        existing_extra = dict(candidate.extra or {})
        extra_changed = False
        if "languages" not in existing_extra and parsed_data.get("languages"):
            existing_extra["languages"] = parsed_data["languages"]
            extra_changed = True
        if extra_changed:
            candidate.extra = existing_extra

        await session.flush()
        logger.info(f"Successfully updated candidate {candidate_id} from parsed resume")

    except Exception as e:
        logger.warning(f"Resume parsing failed for {filename}: {e}")
        # Don't raise - we don't want to block file upload
        return