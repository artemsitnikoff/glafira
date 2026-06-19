import re
from datetime import datetime
from io import BytesIO
from uuid import UUID
from urllib.parse import quote
from xml.sax.saxutils import escape as _xml_escape

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.colors import black, grey, HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether
from reportlab.platypus.flowables import HRFlowable
from reportlab.pdfgen import canvas
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document

from ..models.candidate import Candidate
from ..core.errors import NotFoundError


class PageNumCanvas(canvas.Canvas):
    """Custom canvas для футера на каждой странице"""

    def __init__(self, *args, **kwargs):
        self.candidate_name = kwargs.pop('candidate_name', '')
        # Имя зарегистрированного шрифта (DejaVu на VPS / Helvetica fallback локально)
        self.footer_font = kwargs.pop('footer_font', 'Helvetica')
        super().__init__(*args, **kwargs)

    def showPage(self):
        # Добавляем футер перед отображением страницы
        self._add_page_footer()
        super().showPage()

    def _add_page_footer(self):
        self.saveState()
        font_name = self.footer_font

        self.setFont(font_name, 9)
        self.setFillGray(0.5)

        # Дата в нужном формате
        now = datetime.now()
        months = [
            'января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
            'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря'
        ]
        date_str = f"{now.day} {months[now.month - 1]} {now.year}"

        footer_text = f"{self.candidate_name} • Экспортировано из Глафиры {date_str}"

        # Центрируем футер
        text_width = self.stringWidth(footer_text, font_name, 9)
        x = (A4[0] - text_width) / 2
        y = 1.5 * cm

        self.drawString(x, y, footer_text)
        self.restoreState()


def _register_fonts():
    """Регистрирует шрифты DejaVu для кириллицы с fallback на Helvetica"""
    try:
        pdfmetrics.registerFont(TTFont('DejaVu', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'))
        pdfmetrics.registerFont(TTFont('DejaVu-Bold', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'))
        return 'DejaVu', 'DejaVu-Bold'
    except Exception:
        # Fallback на системные шрифты для локальной среды/тестов
        return 'Helvetica', 'Helvetica-Bold'


def _strip_html(text: str) -> str:
    """Убирает HTML-теги из текста"""
    if not text:
        return ""
    return re.sub(r'<[^>]+>', '', text)


def _safe(text) -> str:
    """Снимает HTML-теги И экранирует XML-спецсимволы (& < >) — иначе reportlab
    Paragraph (парсит мини-HTML) падает на реальных данных вроде «R&D», «M&A»."""
    return _xml_escape(_strip_html(text or ""))


def _full_name(candidate: Candidate) -> str:
    """Формирует полное ФИО кандидата"""
    parts = [candidate.last_name, candidate.first_name]
    if candidate.middle_name:
        parts.append(candidate.middle_name)
    return " ".join(filter(None, parts))


async def load_candidate_for_export(
    session: AsyncSession,
    company_id: UUID,
    candidate_id: UUID
) -> Candidate:
    """Загружает кандидата с relations для экспорта"""
    query = (
        select(Candidate)
        .where(
            Candidate.id == candidate_id,
            Candidate.company_id == company_id,
            Candidate.deleted_at.is_(None)
        )
        .options(
            selectinload(Candidate.experience),
            selectinload(Candidate.skills),
            selectinload(Candidate.education)
        )
    )

    result = await session.execute(query)
    candidate = result.scalar_one_or_none()

    if not candidate:
        raise NotFoundError("Кандидат")

    # Сортируем relations по order_index
    if candidate.experience:
        candidate.experience.sort(key=lambda x: x.order_index)
    if candidate.skills:
        candidate.skills.sort(key=lambda x: x.order_index)
    if candidate.education:
        candidate.education.sort(key=lambda x: x.order_index)

    return candidate


def _format_req_match_item(item) -> str:
    """Форматирует элемент requirements_match — строка или dict."""
    if isinstance(item, str):
        return item
    if isinstance(item, dict):
        parts = []
        if item.get("requirement"):
            parts.append(str(item["requirement"]))
        if item.get("matched") is not None:
            parts.append("✓" if item["matched"] else "✗")
        if item.get("comment"):
            parts.append(str(item["comment"]))
        return " ".join(parts) if parts else str(item)
    return str(item)


_VERDICT_MAP = {
    "good": "рекомендуется",
    "partial": "частично подходит",
    "bad": "не подходит",
}


def build_resume_pdf(candidate: Candidate, ai_analysis: dict | None = None) -> bytes:
    """Создает PDF-резюме кандидата в стиле hh"""
    buffer = BytesIO()
    font_normal, font_bold = _register_fonts()

    # Создаем документ с кастомным canvas для футера
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2.5*cm,
        canvasmaker=lambda *args, **kwargs: PageNumCanvas(
            *args, candidate_name=_full_name(candidate), footer_font=font_normal, **kwargs
        )
    )

    # Стили
    styles = getSampleStyleSheet()

    # Кастомные стили
    header_style = ParagraphStyle(
        'CustomHeader',
        parent=styles['Heading1'],
        fontName=font_bold,
        fontSize=24,
        spaceAfter=6,
        textColor=black
    )

    brand_style = ParagraphStyle(
        'Brand',
        parent=styles['Normal'],
        fontName=font_bold,
        fontSize=14,
        textColor=HexColor('#7E5CF0'),
        alignment=TA_RIGHT
    )

    section_header_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Heading2'],
        fontName=font_bold,
        fontSize=12,
        textColor=grey,
        spaceAfter=8,
        spaceBefore=16,
        textTransform='uppercase'
    )

    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontName=font_normal,
        fontSize=11,
        leading=14
    )

    bold_style = ParagraphStyle(
        'CustomBold',
        parent=styles['Normal'],
        fontName=font_bold,
        fontSize=11,
        leading=14
    )

    contact_style = ParagraphStyle(
        'Contact',
        parent=styles['Normal'],
        fontName=font_normal,
        fontSize=10,
        textColor=grey,
        spaceAfter=4
    )

    # Контент документа
    story = []

    # Шапка с ФИО и брендом
    header_table_data = [
        [
            Paragraph(_safe(_full_name(candidate)), header_style),
            Paragraph("Глафира", brand_style)
        ]
    ]

    header_table = Table(header_table_data, colWidths=[12*cm, 6*cm])
    header_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))

    story.append(header_table)
    story.append(Spacer(1, 8))

    # Контактная информация
    contact_info = []

    if candidate.gender:
        gender_map = {'male': 'Мужчина', 'female': 'Женщина'}
        contact_info.append(gender_map.get(candidate.gender, candidate.gender))

    if candidate.phone:
        contact_info.append(candidate.phone)

    if candidate.email:
        contact_info.append(candidate.email)

    if candidate.city:
        location = f"Проживает: {candidate.city}"
        if candidate.region and candidate.region != candidate.city:
            location += f", {candidate.region}"
        contact_info.append(location)

    for info in contact_info:
        story.append(Paragraph(_safe(info), contact_style))

    story.append(Spacer(1, 20))

    # Секция "Желаемая должность и зарплата"
    if candidate.last_position or candidate.salary_expectation:
        story.append(Paragraph("ЖЕЛАЕМАЯ ДОЛЖНОСТЬ И ЗАРПЛАТА", section_header_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=grey, spaceAfter=8))

        if candidate.last_position:
            story.append(Paragraph(_safe(candidate.last_position), bold_style))

        if candidate.salary_expectation:
            salary_text = f"{candidate.salary_expectation:,} {candidate.currency}".replace(',', ' ')
            story.append(Paragraph(_safe(salary_text), normal_style))

    # Секция "Опыт работы"
    if candidate.experience:
        story.append(Paragraph("ОПЫТ РАБОТЫ", section_header_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=grey, spaceAfter=12))

        for exp in candidate.experience:
            exp_data = []

            # Период работы
            period = exp.period or ""

            # Компания и должность
            company_position = []
            if exp.company:
                company_position.append(Paragraph(_safe(exp.company), bold_style))
            company_position.append(Paragraph(_safe(exp.position), normal_style))

            # Описание
            if exp.description:
                # Разбиваем описание по абзацам
                desc_paragraphs = _strip_html(exp.description).split('\n')
                for para in desc_paragraphs:
                    if para.strip():
                        company_position.append(Paragraph(_safe(para.strip()), normal_style))

            exp_data.append([
                Paragraph(_safe(period), normal_style),
                company_position
            ])

            exp_table = Table(exp_data, colWidths=[3*cm, 13*cm])
            exp_table.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                ('TOPPADDING', (0, 0), (-1, -1), 0),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ]))

            story.append(KeepTogether(exp_table))

    # Секция "Образование"
    if candidate.education:
        story.append(Paragraph("ОБРАЗОВАНИЕ", section_header_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=grey, spaceAfter=8))

        for edu in candidate.education:
            if edu.institution:
                story.append(Paragraph(_safe(edu.institution), bold_style))

            edu_details = []
            if edu.specialty:
                edu_details.append(edu.specialty)
            if edu.years:
                edu_details.append(edu.years)

            if edu_details:
                story.append(Paragraph(_safe(" • ".join(edu_details)), normal_style))

            story.append(Spacer(1, 8))

    # Секция "Навыки"
    skills_text = []
    if candidate.skills:
        skills_text.extend([skill.skill for skill in candidate.skills])

    # Добавляем языки из extra
    if candidate.extra and 'languages' in candidate.extra:
        languages = candidate.extra['languages']
        if isinstance(languages, list):
            for lang in languages:
                if isinstance(lang, str):
                    skills_text.append(lang)
                elif isinstance(lang, dict) and 'name' in lang:
                    skills_text.append(lang['name'])

    if skills_text:
        story.append(Paragraph("НАВЫКИ", section_header_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=grey, spaceAfter=8))
        story.append(Paragraph(_safe(", ".join(skills_text)), normal_style))

    # Секция "Обо мне"
    if candidate.resume_summary:
        story.append(Paragraph("ОБО МНЕ", section_header_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=grey, spaceAfter=8))
        story.append(Paragraph(_safe(candidate.resume_summary), normal_style))

    # Секция «Разбор Глафиры (AI)»
    if ai_analysis:
        story.append(Paragraph("РАЗБОР ГЛАФИРЫ (AI)", section_header_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=grey, spaceAfter=8))

        score = ai_analysis.get("score")
        verdict_raw = ai_analysis.get("verdict", "")
        verdict_ru = _VERDICT_MAP.get(verdict_raw, verdict_raw)
        if score is not None:
            score_line = f"Оценка: {score}/100 — {verdict_ru}"
            story.append(Paragraph(_safe(score_line), bold_style))

        summary = ai_analysis.get("summary", "")
        if summary:
            story.append(Paragraph(_safe(summary), normal_style))

        strengths = ai_analysis.get("strengths") or []
        if strengths:
            story.append(Paragraph(_safe("Сильные стороны:"), bold_style))
            for s in strengths:
                story.append(Paragraph(f"• {_safe(str(s))}", normal_style))

        risks = ai_analysis.get("risks") or []
        if risks:
            story.append(Paragraph(_safe("Риски:"), bold_style))
            for r in risks:
                story.append(Paragraph(f"• {_safe(str(r))}", normal_style))

        req_match = ai_analysis.get("requirements_match") or []
        if req_match:
            story.append(Paragraph(_safe("Соответствие требованиям:"), bold_style))
            for item in req_match:
                story.append(Paragraph(f"• {_safe(_format_req_match_item(item))}", normal_style))

        forecast = ai_analysis.get("forecast", "")
        if forecast:
            story.append(Paragraph(_safe("Прогноз:"), bold_style))
            story.append(Paragraph(_safe(forecast), normal_style))

    # Генерируем PDF
    doc.build(story)
    buffer.seek(0)

    return buffer.read()


def build_resume_docx(candidate: Candidate, ai_analysis: dict | None = None) -> bytes:
    """Создает DOCX-резюме кандидата"""
    doc = Document()

    # Заголовок
    heading = doc.add_heading(_full_name(candidate), 0)

    # Брендинг
    brand_para = doc.add_paragraph("Глафира")
    brand_para.alignment = 2  # Right alignment

    # Контактная информация
    contact_info = []

    if candidate.gender:
        gender_map = {'male': 'Мужчина', 'female': 'Женщина'}
        contact_info.append(gender_map.get(candidate.gender, candidate.gender))

    if candidate.phone:
        contact_info.append(candidate.phone)

    if candidate.email:
        contact_info.append(candidate.email)

    if candidate.city:
        location = f"Проживает: {candidate.city}"
        if candidate.region and candidate.region != candidate.city:
            location += f", {candidate.region}"
        contact_info.append(location)

    for info in contact_info:
        doc.add_paragraph(info)

    # Секция "Желаемая должность и зарплата"
    if candidate.last_position or candidate.salary_expectation:
        doc.add_heading("Желаемая должность и зарплата", level=2)

        if candidate.last_position:
            para = doc.add_paragraph()
            run = para.add_run(candidate.last_position)
            run.bold = True

        if candidate.salary_expectation:
            salary_text = f"{candidate.salary_expectation:,} {candidate.currency}".replace(',', ' ')
            doc.add_paragraph(salary_text)

    # Секция "Опыт работы"
    if candidate.experience:
        doc.add_heading("Опыт работы", level=2)

        for exp in candidate.experience:
            # Период
            if exp.period:
                doc.add_paragraph(exp.period)

            # Компания
            if exp.company:
                para = doc.add_paragraph()
                run = para.add_run(exp.company)
                run.bold = True

            # Должность
            doc.add_paragraph(exp.position)

            # Описание
            if exp.description:
                desc_paragraphs = _strip_html(exp.description).split('\n')
                for para_text in desc_paragraphs:
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())

            doc.add_paragraph()  # Пробел между записями

    # Секция "Образование"
    if candidate.education:
        doc.add_heading("Образование", level=2)

        for edu in candidate.education:
            if edu.institution:
                para = doc.add_paragraph()
                run = para.add_run(edu.institution)
                run.bold = True

            edu_details = []
            if edu.specialty:
                edu_details.append(edu.specialty)
            if edu.years:
                edu_details.append(edu.years)

            if edu_details:
                doc.add_paragraph(" • ".join(edu_details))

    # Секция "Навыки"
    skills_text = []
    if candidate.skills:
        skills_text.extend([skill.skill for skill in candidate.skills])

    # Добавляем языки из extra
    if candidate.extra and 'languages' in candidate.extra:
        languages = candidate.extra['languages']
        if isinstance(languages, list):
            for lang in languages:
                if isinstance(lang, str):
                    skills_text.append(lang)
                elif isinstance(lang, dict) and 'name' in lang:
                    skills_text.append(lang['name'])

    if skills_text:
        doc.add_heading("Навыки", level=2)
        doc.add_paragraph(", ".join(skills_text))

    # Секция "Обо мне"
    if candidate.resume_summary:
        doc.add_heading("Обо мне", level=2)
        doc.add_paragraph(_strip_html(candidate.resume_summary))

    # Секция «Разбор Глафиры (AI)»
    if ai_analysis:
        doc.add_heading("Разбор Глафиры (AI)", level=2)

        score = ai_analysis.get("score")
        verdict_raw = ai_analysis.get("verdict", "")
        verdict_ru = _VERDICT_MAP.get(verdict_raw, verdict_raw)
        if score is not None:
            score_para = doc.add_paragraph()
            score_run = score_para.add_run(f"Оценка: {score}/100 — {verdict_ru}")
            score_run.bold = True

        summary = ai_analysis.get("summary", "")
        if summary:
            doc.add_paragraph(_strip_html(summary))

        strengths = ai_analysis.get("strengths") or []
        if strengths:
            strong_para = doc.add_paragraph()
            strong_para.add_run("Сильные стороны:").bold = True
            for s in strengths:
                doc.add_paragraph(f"• {_strip_html(str(s))}")

        risks = ai_analysis.get("risks") or []
        if risks:
            risks_para = doc.add_paragraph()
            risks_para.add_run("Риски:").bold = True
            for r in risks:
                doc.add_paragraph(f"• {_strip_html(str(r))}")

        req_match = ai_analysis.get("requirements_match") or []
        if req_match:
            req_para = doc.add_paragraph()
            req_para.add_run("Соответствие требованиям:").bold = True
            for item in req_match:
                doc.add_paragraph(f"• {_strip_html(_format_req_match_item(item))}")

        forecast = ai_analysis.get("forecast", "")
        if forecast:
            fc_para = doc.add_paragraph()
            fc_para.add_run("Прогноз:").bold = True
            doc.add_paragraph(_strip_html(forecast))

    # Сохраняем в буфер
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.read()